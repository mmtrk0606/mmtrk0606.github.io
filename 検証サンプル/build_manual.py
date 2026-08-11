# -*- coding: utf-8 -*-
"""
ポートフォリオ用 操作マニュアルサンプル
架空の会議室予約システム「Rezerva」の利用者向けマニュアルを作成した想定
※実在しないサービスです

画面キャプチャは manual_shots/ に置いたモック画面を使用する。
配色は対象システム（Rezerva）のブランドカラーに合わせている。
"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

SHOTS = "manual_shots"
OUT = "out/操作マニュアル_Rezerva_サンプル.docx"

# ── 配色（対象システムのブランドカラー） ────────────────
BRAND = RGBColor(0x1D, 0x4E, 0xD8)
INK = RGBColor(0x11, 0x18, 0x27)
INK2 = RGBColor(0x4B, 0x55, 0x63)
INK3 = RGBColor(0x9C, 0xA3, 0xAF)
RED = RGBColor(0x99, 0x1B, 0x1B)

FONT = "游ゴシック"
BODY_W_MM = 170  # 本文の幅（A4 210mm − 左右20mm）


# ══ 低レベルのヘルパ ═══════════════════════════════════
def set_font(run, size=10.5, bold=False, color=INK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    # 日本語フォントは eastAsia にも指定しないと反映されない
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def shade(el, hexcolor):
    """段落・セルに背景色を敷く"""
    pr = el.get_or_add_tcPr() if el.tag.endswith("tc") else el.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    pr.append(sh)


def cell_border(cell, color="D1D5DB", sz=4, edges=("top", "left", "bottom", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for e in edges:
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def keep_together(row):
    """行が改ページで分断されないようにする（PDF・Pages対策）"""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    trPr.append(el)


def para(doc, text="", size=10.5, bold=False, color=INK, align=None,
         before=0, after=6, indent=0, line=1.6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if indent:
        p.paragraph_format.left_indent = Mm(indent)
    if align is not None:
        p.alignment = align
    if text:
        set_font(p.add_run(text), size, bold, color)
    return p


# ══ 構成要素 ═══════════════════════════════════════════
def h1(doc, num, text):
    """章見出し：番号を色付きで、下に罫線"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(f"{num}　"), 16, True, BRAND)
    set_font(p.add_run(text), 16, True, INK)
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "1D4ED8")
    bd.append(bottom)
    pPr.append(bd)
    return p


def h2(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(f"{num}　{text}"), 12.5, True, INK)
    return p


def step(doc, n, text):
    """番号付きの操作手順"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.first_line_indent = Mm(-8)
    p.paragraph_format.line_spacing = 1.6
    set_font(p.add_run(f"{n}. "), 10.5, True, BRAND)
    set_font(p.add_run(text), 10.5)
    return p


def bullet(doc, text, indent=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Mm(indent + 4)
    p.paragraph_format.first_line_indent = Mm(-4)
    p.paragraph_format.line_spacing = 1.6
    set_font(p.add_run("・"), 10.5, color=INK3)
    set_font(p.add_run(text), 10.5)
    return p


def figure(doc, filename, caption, width_mm=None):
    """画面キャプチャを枠付きで挿入し、下に図番号つきのキャプションを置く"""
    path = os.path.join(SHOTS, filename)
    w, h = Image.open(path).size
    width_mm = width_mm or BODY_W_MM
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell_border(cell, "D1D5DB", 4)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell.paragraphs[0].add_run().add_picture(path, width=Mm(width_mm - 3))
    keep_together(t.rows[0])
    cap = para(doc, caption, 9, color=INK2, align=WD_ALIGN_PARAGRAPH.CENTER,
               before=4, after=14)
    return t


def notice(doc, title, lines, kind="info"):
    """注意・補足のボックス。改ページで割れないようにする"""
    bg, fg = ("EFF5FF", BRAND) if kind == "info" else ("FEF2F2", RED)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell._tc, bg)
    cell_border(cell, "D1D5DB" if kind == "info" else "FECACA", 4)
    cell.width = Mm(BODY_W_MM)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    set_font(p.add_run(("ⓘ " if kind == "info" else "！ ") + title), 10, True, fg)
    for i, line in enumerate(lines):
        q = cell.add_paragraph()
        q.paragraph_format.space_after = Pt(2 if i < len(lines) - 1 else 0)
        q.paragraph_format.line_spacing = 1.6
        set_font(q.add_run(line), 10, color=INK)
    keep_together(t.rows[0])
    para(doc, "", after=12)
    return t


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        c.width = Mm(widths[i])
        shade(c._tc, "1D4ED8")
        cell_border(c, "1D4ED8", 4)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        set_font(p.add_run(htxt), 9.5, True, RGBColor(0xFF, 0xFF, 0xFF))
    keep_together(hdr)
    for r in rows:
        row = t.add_row()
        for i, v in enumerate(r):
            c = row.cells[i]
            c.width = Mm(widths[i])
            cell_border(c, "D1D5DB", 4)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.5
            set_font(p.add_run(v), 9.5)
        keep_together(row)
    para(doc, "", after=12)
    return t


def page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_font(run, 9, color=INK3)
    for el, attr in [("w:fldChar", "begin"), ("w:instrText", None), ("w:fldChar", "end")]:
        e = OxmlElement(el)
        if el == "w:fldChar":
            e.set(qn("w:fldCharType"), attr)
        else:
            e.set(qn("xml:space"), "preserve")
            e.text = " PAGE "
        run._r.append(e)


# ══ 本体 ═══════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin = Mm(24)
sec.bottom_margin = Mm(20)
page_number_footer(sec)

st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

# ── 表紙 ───────────────────────────────────────────────
para(doc, "", after=90)
para(doc, "Rezerva", 34, True, BRAND, WD_ALIGN_PARAGRAPH.CENTER, after=4)
para(doc, "会議室予約システム", 12, color=INK2, align=WD_ALIGN_PARAGRAPH.CENTER, after=52)
para(doc, "操作マニュアル", 22, True, INK, WD_ALIGN_PARAGRAPH.CENTER, after=8)
para(doc, "利用者向け", 12, color=INK2, align=WD_ALIGN_PARAGRAPH.CENTER, after=76)
para(doc, "第 1.0 版", 10.5, color=INK2, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, "2026年8月", 10.5, color=INK2, align=WD_ALIGN_PARAGRAPH.CENTER, after=40)
para(doc, "本書はポートフォリオ用の制作サンプルです。", 9, True, RED,
     WD_ALIGN_PARAGRAPH.CENTER, after=2)
para(doc, "Rezerva は実在しないサービスであり、記載の画面・仕様はすべて架空のものです。",
     9, color=INK3, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ── 目次 ───────────────────────────────────────────────
para(doc, "目次", 16, True, INK, after=14)
TOC = [
    ("1", "はじめに", ["1-1　本書について", "1-2　動作環境", "1-3　用語"]),
    ("2", "ログインとログアウト", ["2-1　ログインする", "2-2　ログアウトする"]),
    ("3", "画面の見方", ["3-1　画面の構成", "3-2　メニューの一覧"]),
    ("4", "会議室を予約する", ["4-1　予約を登録する", "4-2　入力項目", "4-3　予約できないとき"]),
    ("5", "予約を確認・変更する", ["5-1　自分の予約を確認する", "5-2　予約の内容を変更する"]),
    ("6", "予約をキャンセルする", ["6-1　キャンセルの手順", "6-2　キャンセルできる期限"]),
    ("7", "管理者の機能", ["7-1　会議室を追加する", "7-2　会議室を削除する"]),
    ("8", "困ったときは", []),
]
for num, title, subs in TOC:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{num}　"), 11.5, True, BRAND)
    set_font(p.add_run(title), 11.5, True, INK)
    for s in subs:
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Mm(8)
        q.paragraph_format.space_after = Pt(1)
        set_font(q.add_run(s), 10, color=INK2)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ── 1. はじめに ────────────────────────────────────────
h1(doc, "1", "はじめに")

h2(doc, "1-1", "本書について")
para(doc, "本書は、会議室予約システム「Rezerva」を利用する方に向けた操作マニュアルです。"
          "会議室の予約・変更・キャンセルの手順を、実際の画面とあわせて説明します。")
para(doc, "管理者の方のみが使う機能は、第7章にまとめています。")

h2(doc, "1-2", "動作環境")
table(doc,
      ["項目", "内容"],
      [["対応ブラウザ", "Google Chrome ／ Microsoft Edge ／ Safari（いずれも最新版）"],
       ["画面サイズ", "パソコン・タブレット・スマートフォンに対応"],
       ["営業時間", "8:00 〜 22:00（この範囲で予約できます）"],
       ["予約の単位", "30分単位"],
       ["予約できる件数", "1人あたり5件まで（キャンセル済みの予約は含みません）"]],
      [40, 130])

h2(doc, "1-3", "用語")
table(doc,
      ["用語", "意味"],
      [["確定", "予約が成立している状態です。会議室はこの時間帯、他の人が予約できません。"],
       ["キャンセル済", "取り消した予約です。記録は残りますが、会議室は空き枠に戻ります。"],
       ["予約番号", "予約ごとに発行される番号です（例：RZ-20260818-0043）。"],
       ["管理者", "会議室や利用者を管理できる権限です。通常の利用者には表示されません。"]],
      [32, 138])

# ── 2. ログインとログアウト ────────────────────────────
h1(doc, "2", "ログインとログアウト")

h2(doc, "2-1", "ログインする")
step(doc, 1, "ブラウザで、社内から案内された Rezerva のURLを開きます。")
step(doc, 2, "メールアドレスとパスワードを入力します。")
step(doc, 3, "［ログイン］を押します。")
para(doc, "", after=8)
figure(doc, "m-01_ログイン.png", "図 2-1　ログイン画面", width_mm=120)
notice(doc, "パスワードが分からないとき", [
    "パスワードの再発行は、社内のシステム管理者へご依頼ください。",
    "5回続けて誤ると、30分間ログインできなくなります。",
])

h2(doc, "2-2", "ログアウトする")
para(doc, "画面左下の利用者名を押し、［ログアウト］を選びます。"
          "共有のパソコンを使っている場合は、作業が終わったら必ずログアウトしてください。")

# ── 3. 画面の見方 ──────────────────────────────────────
h1(doc, "3", "画面の見方")

h2(doc, "3-1", "画面の構成")
para(doc, "ログインすると、その日の予約一覧が表示されます。"
          "画面の左側がメニュー、右側が選んだメニューの内容です。")
figure(doc, "m-02_予約一覧.png", "図 3-1　予約一覧（ログイン直後の画面）")

h2(doc, "3-2", "メニューの一覧")
table(doc,
      ["メニュー", "できること"],
      [["予約一覧", "選んだ日の予約を、すべての利用者ぶんまとめて確認します。"],
       ["新規予約", "会議室を予約します。"],
       ["自分の予約", "自分が登録した予約の確認・変更・キャンセルを行います。"],
       ["会議室管理", "会議室の追加・変更・削除を行います（管理者のみ）。"],
       ["利用者管理", "利用者の追加・権限の変更を行います（管理者のみ）。"]],
      [34, 136])

# ── 4. 会議室を予約する ────────────────────────────────
h1(doc, "4", "会議室を予約する")

h2(doc, "4-1", "予約を登録する")
step(doc, 1, "メニューから［新規予約］を選びます。")
step(doc, 2, "会議室・利用日・時間・件名を入力します。")
step(doc, 3, "必要に応じて備考を入力します（持ち込む機材や参加人数など）。")
step(doc, 4, "［この内容で予約する］を押します。")
para(doc, "", after=8)
figure(doc, "m-03_新規予約.png", "図 4-1　新規予約の入力画面")
para(doc, "登録すると予約一覧に戻り、画面の上部に予約番号が表示されます。")
figure(doc, "m-04_登録完了.png", "図 4-2　登録が完了した状態（登録した予約が一覧に追加される）")

h2(doc, "4-2", "入力項目")
table(doc,
      ["項目", "必須", "説明"],
      [["会議室", "必須", "予約する会議室を選びます。定員も表示されます。"],
       ["利用日", "必須", "当日以降の日付を選びます。過去の日付は指定できません。"],
       ["時間", "必須", "開始と終了を30分単位で選びます（8:00〜22:00）。"],
       ["件名", "必須", "50文字以内。予約一覧に表示され、他の利用者からも見えます。"],
       ["備考", "任意", "500文字以内。持ち込む機材や連絡事項を記載します。"]],
      [26, 18, 126])

h2(doc, "4-3", "予約できないとき")
para(doc, "次の場合は、登録時にメッセージが表示されて予約できません。")
bullet(doc, "同じ会議室・同じ時間帯に、すでに他の予約が入っている")
bullet(doc, "終了時刻が開始時刻と同じか、それより前になっている")
bullet(doc, "営業時間（8:00〜22:00）の外を指定している")
bullet(doc, "有効な予約をすでに5件持っている")
para(doc, "", after=8)
notice(doc, "時間帯が重なる予約はできません", [
    "たとえば 10:00〜12:00 の予約がある会議室に、11:00〜13:00 の予約は登録できません。",
    "12:00〜13:00 のように、終了と開始が接するだけであれば登録できます。",
], kind="warn")

# ── 5. 予約を確認・変更する ────────────────────────────
h1(doc, "5", "予約を確認・変更する")

h2(doc, "5-1", "自分の予約を確認する")
para(doc, "メニューから［自分の予約］を選ぶと、自分が登録した今後の予約が一覧で表示されます。")
figure(doc, "m-05_自分の予約.png", "図 5-1　自分の予約")

h2(doc, "5-2", "予約の内容を変更する")
step(doc, 1, "［自分の予約］で、変更したい予約の［編集］を押します。")
step(doc, 2, "会議室・利用日・時間・件名を書き換えます。")
step(doc, 3, "［変更を保存］を押します。")
para(doc, "", after=8)
figure(doc, "m-06_予約の編集.png", "図 5-2　予約の編集画面")
notice(doc, "変更できるのは自分の予約だけです", [
    "他の利用者が登録した予約は、一覧で確認はできますが、変更・削除はできません。",
    "他の人の予約を動かす必要がある場合は、予約者ご本人か管理者へご連絡ください。",
])

# ── 6. 予約をキャンセルする ────────────────────────────
h1(doc, "6", "予約をキャンセルする")

h2(doc, "6-1", "キャンセルの手順")
step(doc, 1, "メニューから［自分の予約］を選びます。")
step(doc, 2, "取り消したい予約の［キャンセル］を押します。")
step(doc, 3, "確認のメッセージで［はい］を選びます。")
para(doc, "状態が「キャンセル済」に変わり、その時間帯は他の利用者が予約できるようになります。")

h2(doc, "6-2", "キャンセルできる期限")
notice(doc, "キャンセルは利用開始の1時間前までです", [
    "開始1時間を切った予約はキャンセルできません。［キャンセル］は押せない状態になります。",
    "やむを得ない場合は、会議室の管理担当者へ直接ご連絡ください。",
], kind="warn")

# ── 7. 管理者の機能 ────────────────────────────────────
h1(doc, "7", "管理者の機能")
para(doc, "この章の操作は、管理者の権限を持つ方のみが行えます。"
          "一般の利用者には、メニューの「管理」セクションは表示されません。")

h2(doc, "7-1", "会議室を追加する")
step(doc, 1, "メニューの「管理」から［会議室管理］を選びます。")
step(doc, 2, "［＋ 会議室を追加］を押します。")
step(doc, 3, "名称・定員・設備を入力して保存します。")
para(doc, "追加した会議室は、すぐに［新規予約］の選択肢に表示されます。")
figure(doc, "m-07_会議室管理.png", "図 7-1　会議室管理")

h2(doc, "7-2", "会議室を削除する")
para(doc, "会議室の行の［削除］を押すと、その会議室を一覧から取り除けます。")
notice(doc, "予約が入っている会議室は削除しないでください", [
    "予約が残ったまま削除すると、その予約は会議室名が表示されない状態になり、"
    "予約者本人も取り消せなくなります。",
    "先に予約者へ連絡し、予約がすべてなくなってから削除してください。",
], kind="warn")

# ── 8. 困ったときは ────────────────────────────────────
h1(doc, "8", "困ったときは")
table(doc,
      ["症状", "確認していただきたいこと"],
      [["ログインできない",
        "メールアドレスの入力に誤りがないかご確認ください。"
        "5回続けて誤ると30分間ログインできません。"],
       ["予約が登録できない",
        "「4-3 予約できないとき」をご確認ください。"
        "画面上部のメッセージに原因が表示されます。"],
       ["［キャンセル］が押せない",
        "利用開始まで1時間を切っている可能性があります。"
        "「6-2 キャンセルできる期限」をご確認ください。"],
       ["他の人の予約を消したい",
        "他の利用者の予約は変更・削除できません。"
        "予約者ご本人か管理者へご連絡ください。"],
       ["会議室が選択肢に出てこない",
        "管理者によって削除されている可能性があります。"
        "システム管理者へお問い合わせください。"],
       ["スマートフォンで表が見切れる",
        "表を横方向にスワイプすると、隠れている列を確認できます。"]],
      [42, 128])

para(doc, "", after=20)
para(doc, "本書はポートフォリオ用の制作サンプルです。", 9, True, RED, after=2)
para(doc, "Rezerva は実在しないサービスであり、記載の画面・仕様・数値はすべて架空のものです。",
     9, color=INK3, after=2)
para(doc, "実際のご依頼では、対象システムの実画面を撮影したうえで作成いたします。",
     9, color=INK3, after=0)

doc.save(OUT)
print(f"生成完了: {OUT}")
